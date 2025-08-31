"""
Template Parser for processing base CP templates and identifying fillable fields
"""

import re
import os
from pathlib import Path
from typing import Dict, List, Any, Optional
import logging

try:
    import PyPDF2
    import pdfplumber
    from docx import Document
    from pdf2docx import Converter
except ImportError:
    # Handle import errors gracefully
    PyPDF2 = None
    pdfplumber = None
    Document = None
    Converter = None

logger = logging.getLogger(__name__)

class TemplateParser:
    """Parser for extracting structure and fields from CP templates"""
    
    def __init__(self):
        # Common CP field patterns
        self.field_patterns = {
            'vessel_name': [
                r'\[vessel name\]',
                r'\[vessel\]',
                r'___+\s*vessel',
                r'm\.?v\.?\s*___+',
                r'vessel:?\s*___+'
            ],
            'charterer': [
                r'\[charterer\]',
                r'\[charterers?\]',
                r'___+\s*charterer',
                r'charterer:?\s*___+'
            ],
            'owner': [
                r'\[owner\]',
                r'\[owners?\]',
                r'___+\s*owner',
                r'owner:?\s*___+',
                r'disponent:?\s*___+'
            ],
            'cargo': [
                r'\[cargo\]',
                r'\[commodity\]',
                r'___+\s*cargo',
                r'cargo:?\s*___+',
                r'commodity:?\s*___+'
            ],
            'quantity': [
                r'\[quantity\]',
                r'\[tonnage\]',
                r'___+\s*(?:metric\s*)?tons?',
                r'quantity:?\s*___+',
                r'\d+\s*___+\s*(?:metric\s*)?tons?'
            ],
            'load_port': [
                r'\[load(?:ing)?\s*port\]',
                r'\[port\s*of\s*loading\]',
                r'___+\s*(?:port\s*of\s*)?loading',
                r'loading:?\s*___+'
            ],
            'discharge_port': [
                r'\[discharge\s*port\]',
                r'\[port\s*of\s*discharge\]',
                r'___+\s*(?:port\s*of\s*)?discharge',
                r'discharge:?\s*___+'
            ],
            'freight_rate': [
                r'\[freight\s*rate\]',
                r'\[freight\]',
                r'___+\s*per\s*(?:metric\s*)?ton',
                r'freight:?\s*___+',
                r'usd?\s*___+\s*per\s*(?:mt|ton)'
            ],
            'laycan_start': [
                r'\[laycan\s*start\]',
                r'\[laydays?\s*commence\]',
                r'___+\s*lay.*days?',
                r'laydays?:?\s*___+'
            ],
            'laycan_end': [
                r'\[laycan\s*end\]',
                r'\[cancelling\]',
                r'cancelling:?\s*___+',
                r'___+\s*cancelling'
            ],
            'demurrage': [
                r'\[demurrage\]',
                r'___+\s*per\s*day\s*demurrage',
                r'demurrage:?\s*___+',
                r'usd?\s*___+\s*per\s*day(?:\s*demurrage)?'
            ],
            'despatch': [
                r'\[despatch\]',
                r'\[dispatch\]',
                r'___+\s*per\s*day\s*despatch',
                r'despatch:?\s*___+'
            ],
            'laytime': [
                r'\[laytime\]',
                r'\[lay.*time\]',
                r'___+\s*(?:running\s*)?(?:lay.*)?days?',
                r'laytime:?\s*___+'
            ],
            'notice_time': [
                r'\[notice\]',
                r'\[notice\s*time\]',
                r'___+\s*hours?\s*notice',
                r'notice:?\s*___+'
            ]
        }
        
        # Template type identifiers
        self.template_identifiers = {
            'GENCON': [
                'gencon',
                'general cargo',
                'uniform general charter',
                'dry cargo charter'
            ],
            'NYPE': [
                'nype',
                'new york produce exchange',
                'time charter',
                'new york produce'
            ],
            'SHELLTIME': [
                'shelltime',
                'shell time charter',
                'shell international'
            ],
            'ASBATANKVOY': [
                'asbatankvoy',
                'asba tanker',
                'american standard'
            ]
        }
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse a CP template and extract structure and fields"""
        try:
            # Extract text from template
            text = await self._extract_text(file_path)
            if not text:
                raise ValueError("No text could be extracted from the template")
            
            # Identify template type
            template_type = self._identify_template_type(text)
            
            # Extract fillable fields
            fields = self._extract_fields(text)
            
            # Analyze document structure
            structure = self._analyze_structure(text)
            
            # Extract clauses
            clauses = self._extract_clauses(text)
            
            parsed_data = {
                "original_text": text,
                "template_type": template_type,
                "fields": fields,
                "structure": structure,
                "clauses": clauses,
                "file_info": {
                    "filename": os.path.basename(file_path),
                    "file_type": Path(file_path).suffix.lower()
                }
            }
            
            logger.info(f"Successfully parsed template: {template_type}, {len(fields)} fields found")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing template {file_path}: {str(e)}")
            raise
    
    async def _extract_text(self, file_path: str) -> str:
        """Extract text from template file"""
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.txt':
                return self._extract_from_txt(file_path)
            elif file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        # Try pdfplumber first for better formatting preservation
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        if PyPDF2:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except Exception as e:
                logger.error(f"PyPDF2 also failed: {e}")
                raise
        
        raise ImportError("No PDF processing library available")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not Document:
            raise ImportError("python-docx library not available")
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    
    def _identify_template_type(self, text: str) -> str:
        """Identify the type of CP template"""
        text_lower = text.lower()
        
        for template_type, identifiers in self.template_identifiers.items():
            for identifier in identifiers:
                if identifier.lower() in text_lower:
                    return template_type
        
        # Default fallback
        return "UNKNOWN"
    
    def _extract_fields(self, text: str) -> List[Dict[str, Any]]:
        """Extract fillable fields from the template"""
        fields = []
        text_lower = text.lower()
        
        for field_type, patterns in self.field_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, text_lower, re.IGNORECASE | re.MULTILINE)
                for match in matches:
                    field = {
                        "type": field_type,
                        "pattern": pattern,
                        "match": match.group(0),
                        "position": match.span(),
                        "context": self._get_context(text, match.span(), 50),
                        "confidence": 0.8
                    }
                    
                    # Avoid duplicates
                    if not self._is_duplicate_field(fields, field):
                        fields.append(field)
        
        # Sort fields by position
        fields.sort(key=lambda x: x['position'][0])
        
        logger.info(f"Extracted {len(fields)} fields from template")
        return fields
    
    def _is_duplicate_field(self, existing_fields: List[Dict], new_field: Dict) -> bool:
        """Check if a field is a duplicate"""
        for field in existing_fields:
            # Check if positions overlap
            if (field['position'][0] <= new_field['position'][1] and 
                field['position'][1] >= new_field['position'][0]):
                return True
        return False
    
    def _get_context(self, text: str, position: tuple, context_length: int = 50) -> str:
        """Get context around a matched position"""
        start = max(0, position[0] - context_length)
        end = min(len(text), position[1] + context_length)
        return text[start:end].strip()
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze the document structure"""
        lines = text.split('\n')
        
        structure = {
            "total_lines": len(lines),
            "sections": [],
            "headers": [],
            "numbered_clauses": [],
            "bullet_points": []
        }
        
        # Find headers (lines in all caps or with specific patterns)
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Check for headers
            if (line_stripped.isupper() and len(line_stripped) > 5 and 
                not any(char.isdigit() for char in line_stripped)):
                structure["headers"].append({
                    "text": line_stripped,
                    "line_number": i + 1
                })
            
            # Check for numbered clauses
            if re.match(r'^\s*\d+\.', line_stripped):
                structure["numbered_clauses"].append({
                    "text": line_stripped[:100] + "..." if len(line_stripped) > 100 else line_stripped,
                    "line_number": i + 1,
                    "clause_number": re.match(r'^\s*(\d+)\.', line_stripped).group(1)
                })
            
            # Check for bullet points
            if re.match(r'^\s*[-•*]\s+', line_stripped):
                structure["bullet_points"].append({
                    "text": line_stripped[:100] + "..." if len(line_stripped) > 100 else line_stripped,
                    "line_number": i + 1
                })
        
        return structure
    
    def _extract_clauses(self, text: str) -> List[Dict[str, Any]]:
        """Extract main clauses from the template"""
        clauses = []
        
        # Split text into potential clauses
        # Look for numbered sections, paragraph breaks, etc.
        sections = re.split(r'\n\s*\n', text)
        
        for i, section in enumerate(sections):
            section = section.strip()
            if len(section) < 50:  # Skip very short sections
                continue
            
            # Check if it's a numbered clause
            clause_match = re.match(r'^\s*(\d+)\.\s*(.+)', section, re.DOTALL)
            if clause_match:
                clause_number = clause_match.group(1)
                clause_text = clause_match.group(2)
                
                clauses.append({
                    "number": clause_number,
                    "title": self._extract_clause_title(clause_text),
                    "text": clause_text[:500] + "..." if len(clause_text) > 500 else clause_text,
                    "full_text": clause_text,
                    "position": i,
                    "type": "numbered"
                })
            else:
                # Try to identify clause by content
                title = self._extract_clause_title(section)
                if title:
                    clauses.append({
                        "number": None,
                        "title": title,
                        "text": section[:500] + "..." if len(section) > 500 else section,
                        "full_text": section,
                        "position": i,
                        "type": "paragraph"
                    })
        
        logger.info(f"Extracted {len(clauses)} clauses from template")
        return clauses
    
    def _extract_clause_title(self, text: str) -> str:
        """Extract title from clause text"""
        # Take first line or first sentence as title
        lines = text.split('\n')
        first_line = lines[0].strip()
        
        # If first line is short and descriptive, use it
        if len(first_line) < 100 and len(first_line.split()) > 1:
            return first_line
        
        # Otherwise, extract first few words
        words = text.split()
        if len(words) > 10:
            return ' '.join(words[:10]) + "..."
        
        return first_line[:50] + "..." if len(first_line) > 50 else first_line

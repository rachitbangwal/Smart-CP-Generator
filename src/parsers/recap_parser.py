"""
Recap Document Parser for extracting commercial terms from recap documents
"""

import re
import os
from pathlib import Path
from typing import Dict, List, Any, Optional
import logging

import spacy
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

try:
    import PyPDF2
    import pdfplumber
    from docx import Document
except ImportError:
    # Handle import errors gracefully
    PyPDF2 = None
    pdfplumber = None
    Document = None

logger = logging.getLogger(__name__)

class RecapParser:
    """Parser for extracting commercial terms from recap documents"""
    
    def __init__(self):
        self.nlp = None
        self.stop_words = set()
        self._initialize_nlp()
        
        # Commercial terms patterns
        self.term_patterns = {
            'freight': [
                r'freight[:\s]+(\$?[\d,]+\.?\d*)\s*(per|\/)\s*(mt|ton|tonne)',
                r'freight rate[:\s]+(\$?[\d,]+\.?\d*)',
                r'(\$?[\d,]+\.?\d*)\s*(per|\/)\s*(mt|ton|tonne)\s*freight'
            ],
            'laycan': [
                r'laycan[:\s]+(\d{1,2}[-\/]\d{1,2}[-\/]\d{2,4})\s*[-to]\s*(\d{1,2}[-\/]\d{1,2}[-\/]\d{2,4})',
                r'lay.*days?[:\s]+(\d{1,2}[-\/]\d{1,2}[-\/]\d{2,4})\s*[-to]\s*(\d{1,2}[-\/]\d{1,2}[-\/]\d{2,4})',
                r'cancelling[:\s]+(\d{1,2}[-\/]\d{1,2}[-\/]\d{2,4})'
            ],
            'cargo': [
                r'cargo[:\s]+(.+?)(?:\n|$)',
                r'commodity[:\s]+(.+?)(?:\n|$)',
                r'(\d+(?:,\d+)*)\s*(mt|tons?|tonnes?)\s*(?:of\s+)?(.+?)(?:\n|$)'
            ],
            'load_port': [
                r'load(?:ing)?\s*port[:\s]+(.+?)(?:\n|$)',
                r'load(?:ing)?[:\s]+(.+?)(?:\n|$)',
                r'from[:\s]+(.+?)(?:\n|$)'
            ],
            'discharge_port': [
                r'discharge\s*port[:\s]+(.+?)(?:\n|$)',
                r'discharge[:\s]+(.+?)(?:\n|$)',
                r'to[:\s]+(.+?)(?:\n|$)'
            ],
            'demurrage': [
                r'demurrage[:\s]+(\$?[\d,]+\.?\d*)\s*per\s*day',
                r'dem[:\s]+(\$?[\d,]+\.?\d*)\s*\/\s*day',
                r'(\$?[\d,]+\.?\d*)\s*per\s*day\s*demurrage'
            ],
            'despatch': [
                r'despatch[:\s]+(\$?[\d,]+\.?\d*)\s*per\s*day',
                r'desp[:\s]+(\$?[\d,]+\.?\d*)\s*\/\s*day',
                r'(\$?[\d,]+\.?\d*)\s*per\s*day\s*despatch'
            ],
            'quantity': [
                r'quantity[:\s]+(\d+(?:,\d+)*)\s*(mt|tons?|tonnes?)',
                r'(\d+(?:,\d+)*)\s*(mt|tons?|tonnes?)\s*(?:cargo|commodity)',
                r'about\s+(\d+(?:,\d+)*)\s*(mt|tons?|tonnes?)'
            ],
            'vessel': [
                r'vessel[:\s]+(.+?)(?:\n|$)',
                r'ship[:\s]+(.+?)(?:\n|$)',
                r'm[\/]v\s+(.+?)(?:\n|$)'
            ],
            'charterer': [
                r'charterer[:\s]+(.+?)(?:\n|$)',
                r'chtr[:\s]+(.+?)(?:\n|$)'
            ],
            'owner': [
                r'owner[:\s]+(.+?)(?:\n|$)',
                r'disponent[:\s]+(.+?)(?:\n|$)'
            ]
        }
    
    def _initialize_nlp(self):
        """Initialize NLP components"""
        try:
            # Load spaCy model
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("spaCy model loaded successfully")
        except OSError:
            logger.warning("spaCy model not found. NLP features will be limited.")
            self.nlp = None
        
        try:
            # Download NLTK data if needed
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            self.stop_words = set(stopwords.words('english'))
            logger.info("NLTK components loaded successfully")
        except Exception as e:
            logger.warning(f"Error loading NLTK components: {e}")
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse a recap document and extract commercial terms"""
        try:
            # Extract text from file
            text = await self._extract_text(file_path)
            if not text:
                raise ValueError("No text could be extracted from the file")
            
            # Extract commercial terms
            terms = self._extract_terms(text)
            
            # Perform NLP analysis
            nlp_analysis = self._perform_nlp_analysis(text)
            
            # Structure the parsed data
            parsed_data = {
                "original_text": text,
                "terms": terms,
                "nlp_analysis": nlp_analysis,
                "file_info": {
                    "filename": os.path.basename(file_path),
                    "file_type": Path(file_path).suffix.lower()
                }
            }
            
            logger.info(f"Successfully parsed recap document: {len(terms)} terms extracted")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing recap document {file_path}: {str(e)}")
            raise
    
    async def _extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.txt':
                return self._extract_from_txt(file_path)
            elif file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        # Try pdfplumber first (better for tables and complex layouts)
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        if PyPDF2:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except Exception as e:
                logger.error(f"PyPDF2 also failed: {e}")
                raise
        
        raise ImportError("No PDF processing library available")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not Document:
            raise ImportError("python-docx library not available")
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    
    def _extract_terms(self, text: str) -> Dict[str, Any]:
        """Extract commercial terms using regex patterns"""
        terms = {}
        text_lower = text.lower()
        
        for term_type, patterns in self.term_patterns.items():
            matches = []
            
            for pattern in patterns:
                regex_matches = re.finditer(pattern, text_lower, re.IGNORECASE | re.MULTILINE)
                for match in regex_matches:
                    if match.groups():
                        matches.append({
                            "value": match.group(1).strip() if len(match.groups()) >= 1 else match.group(0).strip(),
                            "full_match": match.group(0).strip(),
                            "confidence": 0.8,  # Base confidence for regex matches
                            "position": match.span()
                        })
            
            if matches:
                # Remove duplicates and keep the best match
                unique_matches = self._deduplicate_matches(matches)
                terms[term_type] = unique_matches
        
        return terms
    
    def _deduplicate_matches(self, matches: List[Dict]) -> List[Dict]:
        """Remove duplicate matches and keep the best ones"""
        if not matches:
            return []
        
        # Sort by confidence and position
        matches.sort(key=lambda x: (x['confidence'], -x['position'][0]), reverse=True)
        
        # Remove near-duplicates
        unique_matches = []
        for match in matches:
            is_duplicate = False
            for existing in unique_matches:
                # Check if positions overlap significantly
                overlap = min(match['position'][1], existing['position'][1]) - max(match['position'][0], existing['position'][0])
                if overlap > 0:
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                unique_matches.append(match)
        
        return unique_matches
    
    def _perform_nlp_analysis(self, text: str) -> Dict[str, Any]:
        """Perform NLP analysis on the text"""
        analysis = {
            "entities": [],
            "key_phrases": [],
            "sentiment": None,
            "language": "en"
        }
        
        if not self.nlp:
            return analysis
        
        try:
            # Process text with spaCy
            doc = self.nlp(text)
            
            # Extract named entities
            for ent in doc.ents:
                if ent.label_ in ['ORG', 'GPE', 'MONEY', 'DATE', 'PRODUCT']:
                    analysis["entities"].append({
                        "text": ent.text,
                        "label": ent.label_,
                        "confidence": 0.8,
                        "start": ent.start_char,
                        "end": ent.end_char
                    })
            
            # Extract key phrases (noun chunks)
            for chunk in doc.noun_chunks:
                if len(chunk.text.split()) > 1 and chunk.text.lower() not in self.stop_words:
                    analysis["key_phrases"].append({
                        "text": chunk.text,
                        "pos": chunk.root.pos_,
                        "confidence": 0.6
                    })
            
            logger.info(f"NLP analysis completed: {len(analysis['entities'])} entities, {len(analysis['key_phrases'])} key phrases")
            
        except Exception as e:
            logger.warning(f"Error in NLP analysis: {e}")
        
        return analysis

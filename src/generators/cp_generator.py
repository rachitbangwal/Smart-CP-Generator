"""
Charter Party Generator - Core engine for generating filled charter parties
"""

import os
import re
import logging
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_COLOR_INDEX
    import nltk
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
except ImportError:
    # Handle import errors gracefully
    Document = None
    nltk = None
    TfidfVectorizer = None
    cosine_similarity = None
    np = None

logger = logging.getLogger(__name__)

class CPGenerator:
    """Charter Party Generator for creating filled CP documents"""
    
    def __init__(self):
        self.similarity_threshold = 0.3
        self.confidence_threshold = 0.6
        
        # Term mapping rules
        self.term_mappings = {
            'freight': ['freight_rate', 'freight'],
            'cargo': ['cargo', 'commodity'],
            'vessel': ['vessel_name', 'vessel'],
            'charterer': ['charterer', 'chtr'],
            'owner': ['owner', 'disponent'],
            'load_port': ['load_port', 'loading_port', 'port_of_loading'],
            'discharge_port': ['discharge_port', 'discharging_port', 'port_of_discharge'],
            'quantity': ['quantity', 'tonnage'],
            'demurrage': ['demurrage', 'dem'],
            'despatch': ['despatch', 'dispatch', 'desp'],
            'laycan': ['laycan_start', 'laycan_end', 'laycan'],
            'laytime': ['laytime', 'lay_time']
        }
        
        # Initialize NLP components
        self._initialize_nlp()
    
    def _initialize_nlp(self):
        """Initialize NLP components for semantic matching"""
        try:
            if nltk:
                nltk.download('punkt', quiet=True)
                nltk.download('stopwords', quiet=True)
            
            if TfidfVectorizer:
                self.vectorizer = TfidfVectorizer(
                    stop_words='english',
                    max_features=1000,
                    ngram_range=(1, 2)
                )
            else:
                self.vectorizer = None
                
            logger.info("NLP components initialized successfully")
        except Exception as e:
            logger.warning(f"Error initializing NLP components: {e}")
            self.vectorizer = None
    
    async def generate(self, 
                      template_data: Dict[str, Any], 
                      recap_data: Dict[str, Any], 
                      output_format: str = "docx") -> Dict[str, Any]:
        """Generate a filled charter party document"""
        try:
            logger.info("Starting charter party generation")
            
            # Extract terms from recap
            recap_terms = self._extract_recap_terms(recap_data)
            
            # Map recap terms to template fields
            field_mappings = await self._map_terms_to_fields(recap_terms, template_data)
            
            # Generate the filled document
            filled_document = await self._fill_template(template_data, field_mappings, output_format)
            
            # Track changes
            changes = self._track_changes(template_data, field_mappings)
            
            # Validate the generated document
            validation_result = self._validate_generated_document(filled_document, field_mappings)
            
            result = {
                "filled_document": filled_document,
                "changes": changes,
                "field_mappings": field_mappings,
                "validation": validation_result,
                "statistics": {
                    "fields_filled": len([m for m in field_mappings if m.get("filled", False)]),
                    "total_fields": len(field_mappings),
                    "confidence_score": self._calculate_overall_confidence(field_mappings)
                }
            }
            
            logger.info(f"Charter party generation completed: {result['statistics']['fields_filled']} fields filled")
            return result
            
        except Exception as e:
            logger.error(f"Error generating charter party: {str(e)}")
            raise
    
    def _extract_recap_terms(self, recap_data: Dict[str, Any]) -> Dict[str, Any]:
        """Extract and normalize terms from recap data"""
        extracted_terms = {}
        
        # Get parsed terms from recap
        recap_terms = recap_data.get("terms", {})
        
        for term_type, term_data in recap_terms.items():
            if isinstance(term_data, list) and term_data:
                # Take the highest confidence match
                best_match = max(term_data, key=lambda x: x.get("confidence", 0))
                extracted_terms[term_type] = {
                    "value": best_match.get("value", ""),
                    "confidence": best_match.get("confidence", 0.5),
                    "original_match": best_match.get("full_match", ""),
                    "source": "regex_extraction"
                }
        
        # Extract from NLP analysis if available
        nlp_analysis = recap_data.get("nlp_analysis", {})
        entities = nlp_analysis.get("entities", [])
        
        for entity in entities:
            entity_type = self._map_entity_to_term_type(entity.get("label", ""))
            if entity_type and entity_type not in extracted_terms:
                extracted_terms[entity_type] = {
                    "value": entity.get("text", ""),
                    "confidence": entity.get("confidence", 0.6),
                    "original_match": entity.get("text", ""),
                    "source": "nlp_extraction"
                }
        
        logger.info(f"Extracted {len(extracted_terms)} terms from recap")
        return extracted_terms
    
    async def _map_terms_to_fields(self, 
                                  recap_terms: Dict[str, Any], 
                                  template_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Map recap terms to template fields using semantic matching"""
        field_mappings = []
        template_fields = template_data.get("structured_fields", [])
        
        for field in template_fields:
            field_type = field.get("type", "")
            field_context = field.get("context", "")
            
            # Try direct mapping first
            mapped_term = self._find_direct_mapping(field_type, recap_terms)
            
            if not mapped_term:
                # Try semantic mapping
                mapped_term = await self._find_semantic_mapping(field, recap_terms)
            
            mapping = {
                "field_id": field.get("id", ""),
                "field_type": field_type,
                "field_position": field.get("position", (0, 0)),
                "field_context": field_context,
                "mapped_term": mapped_term,
                "filled": mapped_term is not None,
                "confidence": mapped_term.get("confidence", 0.0) if mapped_term else 0.0,
                "mapping_method": mapped_term.get("mapping_method", "none") if mapped_term else "none"
            }
            
            field_mappings.append(mapping)
        
        logger.info(f"Mapped {len([m for m in field_mappings if m['filled']])} out of {len(field_mappings)} fields")
        return field_mappings
    
    def _find_direct_mapping(self, field_type: str, recap_terms: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Find direct mapping between field type and recap terms"""
        # Check for exact match
        if field_type in recap_terms:
            term = recap_terms[field_type].copy()
            term["mapping_method"] = "direct"
            return term
        
        # Check for mapped term types
        for recap_term, mapped_types in self.term_mappings.items():
            if field_type in mapped_types and recap_term in recap_terms:
                term = recap_terms[recap_term].copy()
                term["mapping_method"] = "mapped"
                return term
        
        return None
    
    async def _find_semantic_mapping(self, field: Dict[str, Any], recap_terms: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Find semantic mapping using NLP similarity"""
        if not self.vectorizer:
            return None
        
        field_context = field.get("context", "")
        if not field_context:
            return None
        
        best_match = None
        best_similarity = 0.0
        
        # Create contexts for comparison
        contexts = [field_context]
        term_keys = []
        
        for term_key, term_data in recap_terms.items():
            contexts.append(term_data.get("original_match", ""))
            term_keys.append(term_key)
        
        try:
            # Calculate similarities
            tfidf_matrix = self.vectorizer.fit_transform(contexts)
            similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
            
            # Find best match above threshold
            for i, similarity in enumerate(similarities):
                if similarity > self.similarity_threshold and similarity > best_similarity:
                    best_similarity = similarity
                    best_match = recap_terms[term_keys[i]].copy()
                    best_match["mapping_method"] = "semantic"
                    best_match["similarity_score"] = float(similarity)
            
        except Exception as e:
            logger.warning(f"Error in semantic mapping: {e}")
        
        return best_match
    
    async def _fill_template(self, 
                           template_data: Dict[str, Any], 
                           field_mappings: List[Dict[str, Any]], 
                           output_format: str) -> Dict[str, Any]:
        """Fill the template with mapped values"""
        original_text = template_data.get("original_data", {}).get("original_text", "")
        filled_text = original_text
        
        # Sort mappings by position (reverse order to avoid position shifts)
        sorted_mappings = sorted(
            [m for m in field_mappings if m.get("filled", False)],
            key=lambda x: x.get("field_position", (0, 0))[0],
            reverse=True
        )
        
        modifications = []
        
        for mapping in sorted_mappings:
            position = mapping.get("field_position", (0, 0))
            mapped_term = mapping.get("mapped_term", {})
            value = mapped_term.get("value", "")
            
            if value and position[0] < len(filled_text):
                # Replace the field placeholder with the actual value
                start, end = position
                old_text = filled_text[start:end]
                filled_text = filled_text[:start] + value + filled_text[end:]
                
                modifications.append({
                    "position": position,
                    "old_text": old_text,
                    "new_text": value,
                    "field_type": mapping.get("field_type", ""),
                    "confidence": mapping.get("confidence", 0.0)
                })
        
        # Create output based on format
        if output_format.lower() == "docx":
            filled_document = await self._create_docx_output(filled_text, modifications)
        elif output_format.lower() == "html":
            filled_document = await self._create_html_output(filled_text, modifications)
        else:
            filled_document = {
                "content": filled_text,
                "format": "text",
                "modifications": modifications
            }
        
        return filled_document
    
    async def _create_docx_output(self, filled_text: str, modifications: List[Dict]) -> Dict[str, Any]:
        """Create DOCX output with change tracking"""
        if not Document:
            raise ImportError("python-docx library not available")
        
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run("CHARTER PARTY")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = 1  # Center alignment
        
        # Add generation info
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        info_run.font.size = Pt(10)
        info_run.italic = True
        
        # Add content with modifications highlighted
        content_text = filled_text
        current_pos = 0
        
        for modification in sorted(modifications, key=lambda x: x["position"][0]):
            pos_start, pos_end = modification["position"]
            
            # Add text before modification
            if current_pos < pos_start:
                doc.add_paragraph(content_text[current_pos:pos_start])
            
            # Add modified text with highlighting
            modified_para = doc.add_paragraph()
            modified_run = modified_para.add_run(modification["new_text"])
            modified_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            modified_run.bold = True
            
            current_pos = pos_end
        
        # Add remaining text
        if current_pos < len(content_text):
            doc.add_paragraph(content_text[current_pos:])
        
        return {
            "document": doc,
            "format": "docx",
            "modifications": modifications,
            "content": filled_text
        }
    
    async def _create_html_output(self, filled_text: str, modifications: List[Dict]) -> Dict[str, Any]:
        """Create HTML output with change tracking"""
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Charter Party</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .header {{ text-align: center; font-size: 18px; font-weight: bold; margin-bottom: 20px; }}
                .info {{ font-size: 10px; font-style: italic; margin-bottom: 20px; }}
                .content {{ white-space: pre-wrap; line-height: 1.6; }}
                .modified {{ background-color: yellow; font-weight: bold; }}
                .change-summary {{ margin-top: 30px; border-top: 1px solid #ccc; padding-top: 20px; }}
            </style>
        </head>
        <body>
            <div class="header">CHARTER PARTY</div>
            <div class="info">Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
            <div class="content">{self._highlight_modifications_in_html(filled_text, modifications)}</div>
            
            <div class="change-summary">
                <h3>Changes Made:</h3>
                <ul>
        """
        
        for i, mod in enumerate(modifications, 1):
            html_content += f"""
                <li>Field {i}: {mod['field_type']} - "{mod['old_text']}" → "{mod['new_text']}" (Confidence: {mod['confidence']:.2f})</li>
            """
        
        html_content += """
                </ul>
            </div>
        </body>
        </html>
        """
        
        return {
            "content": html_content,
            "format": "html",
            "modifications": modifications
        }
    
    def _highlight_modifications_in_html(self, text: str, modifications: List[Dict]) -> str:
        """Highlight modifications in HTML"""
        if not modifications:
            return text
        
        # Sort modifications by position
        sorted_mods = sorted(modifications, key=lambda x: x["position"][0])
        
        result = ""
        current_pos = 0
        
        for mod in sorted_mods:
            pos_start, pos_end = mod["position"]
            
            # Add text before modification
            result += text[current_pos:pos_start]
            
            # Add highlighted modification
            result += f'<span class="modified">{mod["new_text"]}</span>'
            
            current_pos = pos_end
        
        # Add remaining text
        result += text[current_pos:]
        
        return result
    
    def _track_changes(self, template_data: Dict[str, Any], field_mappings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Track all changes made to the template"""
        changes = []
        
        for mapping in field_mappings:
            if mapping.get("filled", False):
                mapped_term = mapping.get("mapped_term", {})
                
                change = {
                    "change_id": f"change_{len(changes) + 1}",
                    "field_id": mapping.get("field_id", ""),
                    "field_type": mapping.get("field_type", ""),
                    "position": mapping.get("field_position", (0, 0)),
                    "original_text": mapping.get("field_context", ""),
                    "new_value": mapped_term.get("value", ""),
                    "confidence": mapping.get("confidence", 0.0),
                    "mapping_method": mapped_term.get("mapping_method", ""),
                    "source_term": mapped_term.get("original_match", ""),
                    "timestamp": datetime.now().isoformat()
                }
                
                changes.append(change)
        
        return changes
    
    def _validate_generated_document(self, filled_document: Dict[str, Any], field_mappings: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Validate the generated document"""
        validation = {
            "is_valid": True,
            "errors": [],
            "warnings": [],
            "completeness_score": 0.0,
            "confidence_score": 0.0
        }
        
        filled_fields = [m for m in field_mappings if m.get("filled", False)]
        total_fields = len(field_mappings)
        
        # Calculate completeness
        if total_fields > 0:
            validation["completeness_score"] = len(filled_fields) / total_fields
        
        # Calculate average confidence
        if filled_fields:
            validation["confidence_score"] = sum(m.get("confidence", 0.0) for m in filled_fields) / len(filled_fields)
        
        # Check for critical missing fields
        critical_fields = ["vessel_name", "charterer", "owner", "cargo"]
        missing_critical = []
        
        for field_type in critical_fields:
            if not any(m.get("field_type") == field_type and m.get("filled", False) for m in field_mappings):
                missing_critical.append(field_type)
        
        if missing_critical:
            validation["errors"].append(f"Missing critical fields: {', '.join(missing_critical)}")
            validation["is_valid"] = False
        
        # Check confidence levels
        low_confidence_fields = [m for m in filled_fields if m.get("confidence", 0.0) < self.confidence_threshold]
        if low_confidence_fields:
            validation["warnings"].append(f"{len(low_confidence_fields)} fields have low confidence scores")
        
        # Check completeness
        if validation["completeness_score"] < 0.7:
            validation["warnings"].append("Document may be incomplete - less than 70% of fields filled")
        
        return validation
    
    def _calculate_overall_confidence(self, field_mappings: List[Dict[str, Any]]) -> float:
        """Calculate overall confidence score for the generation"""
        filled_fields = [m for m in field_mappings if m.get("filled", False)]
        
        if not filled_fields:
            return 0.0
        
        return sum(m.get("confidence", 0.0) for m in filled_fields) / len(filled_fields)
    
    def _map_entity_to_term_type(self, entity_label: str) -> Optional[str]:
        """Map NLP entity labels to term types"""
        mapping = {
            "ORG": "charterer",  # Could be charterer or owner
            "GPE": "load_port",  # Geopolitical entity - likely a port
            "MONEY": "freight_rate",
            "DATE": "laycan_start",
            "PRODUCT": "cargo"
        }
        
        return mapping.get(entity_label)
